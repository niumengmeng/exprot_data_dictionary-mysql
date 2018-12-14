#!/usr/bin/env python3
# coding=utf-8
# @Time    : 2018/12/14 上午10:00
# @Author  : 牛萌萌
# @Desc    : 批量导出mysql的数据字典,只需把xxxx改成自己所需的配置就可以
# @license : Copyright(C), ideadata
# @Contact : love_itmeng@163.com
import pymysql
from docx import Document
connection=pymysql.connect(host='xxxx',
                           user='xxxx',
                           password='xxxx',
                           db='information_schema',
                           port=3306,
                           charset='utf8')
schema='xxxx' #导出的数据库的名字
cursor=connection.cursor()
# 取库下所有数据表名字和表注释
sql="select table_name,table_comment from information_schema.tables where TABLE_SCHEMA = '"+str(schema)+"'"
cursor.execute(sql)
tableInfoList=cursor.fetchall()
doc=Document()
# 得到表的字段所需信息
for tableInfo in tableInfoList:
    tableName=tableInfo[0]
    tableComment=tableInfo[1]
    # 表名 + 表注释
    table_explain = tableName+"   表注释："+tableComment
    # 查询表
    tableInfoSql="SELECT C.ORDINAL_POSITION AS ‘No’, C.COLUMN_NAME AS '字段名', C.COLUMN_TYPE AS '数据类型',C.IS_NULLABLE AS '允许为空',C.EXTRA AS 'PK', C.COLUMN_DEFAULT AS '默认值',C.COLUMN_COMMENT AS '注释' FROM information_schema.COLUMNS C INNER JOIN TABLES T ON C.TABLE_SCHEMA = T.TABLE_SCHEMA AND C.TABLE_NAME = T.TABLE_NAME WHERE T.TABLE_SCHEMA = '"+str(schema)+"' and T.TABLE_NAME='"+str(tableName)+"'"
    cursor.execute(tableInfoSql)
    # 取得表字段信息
    tableColumnInfoList = cursor.fetchall()
    # doc文档换行
    p = doc.add_paragraph('')
    # 设置样式
    p.add_run(table_explain, style="Heading 1 Char")
    row=cursor.rowcount
    # doc文档创建table表格
    table=doc.add_table(rows=1,cols=7)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    # hdr_cells 依次设置列数据
    hdr_cells[0].text='No' # 排序 序号
    hdr_cells[1].text='字段名'
    hdr_cells[2].text = '字段类型'
    hdr_cells[3].text = '允许为空'
    hdr_cells[4].text = 'PK'
    hdr_cells[5].text = '默认值'
    hdr_cells[6].text = '注释'
    # 填写字段对应的数据
    for tableColumn in tableColumnInfoList:
        new_cells = table.add_row().cells
        # hdr_cells 依次设置行数据
        new_cells[0].text=str(tableColumn[0])
        new_cells[1].text=tableColumn[1]
        new_cells[2].text=tableColumn[2]
        new_cells[3].text = tableColumn[3]
        new_cells[4].text = tableColumn[4]
        new_cells[5].text = tableColumn[5] if tableColumn[5] else ''
        new_cells[6].text = tableColumn[6]
    p = doc.add_paragraph('')
# 保存doc，可指定路径 切记加上文件后缀
doc.save('data_dictionary.doc')